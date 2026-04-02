# -------------------------------
# Imports
# -------------------------------
import os
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from pydub import AudioSegment
import requests
import azure.cognitiveservices.speech as speechsdk
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

load_dotenv()

# -------------------------------
# Azure Credentials (loaded from .env)
# -------------------------------
SPEECH_KEY = os.environ["SPEECH_KEY"]
SPEECH_REGION = os.environ["SPEECH_REGION"]
OPENAI_KEY = os.environ["OPENAI_KEY"]
OPENAI_ENDPOINT = os.environ["OPENAI_ENDPOINT"]
OPENAI_DEPLOYMENT = os.environ["OPENAI_DEPLOYMENT"]

# -------------------------------
# Helper: Convert audio to WAV
# -------------------------------
def ensure_wav(file_path):
    ext = file_path.split('.')[-1].lower()
    if ext != "wav":
        audio = AudioSegment.from_file(file_path)
        # Export as WAV using the same filename (overwrite)
        audio.export(file_path, format="wav")
        print(f"🎵 Converted {file_path} → {file_path} (overwritten as WAV)")
    return file_path

# -------------------------------
# Transcription
# -------------------------------
def transcribe_audio(file_path):
    file_path = ensure_wav(file_path)

    # Configure Speech
    speech_config = speechsdk.SpeechConfig(subscription=SPEECH_KEY, region=SPEECH_REGION)
    speech_config.speech_recognition_language = "en-US"

    # Set to use fast/short-form transcription
    # In the SDK, this is done via "speech_recognition_language" + default model, or via deployment for custom fast models
    # Using default short-form model here
    audio_config = speechsdk.AudioConfig(filename=file_path)
    recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)

    print("⏳ Starting FAST continuous transcription...")

    full_text = []

    # Event handler for recognized segments
    def handle_recognized(evt):
        if evt.result.reason == speechsdk.ResultReason.RecognizedSpeech:
            full_text.append(evt.result.text)
            print(evt.result.text)  # optional live output

    recognizer.recognized.connect(handle_recognized)

    done = False
    def stop_cb(evt):
        nonlocal done
        done = True

    recognizer.session_stopped.connect(stop_cb)
    recognizer.canceled.connect(stop_cb)

    # Start continuous recognition
    recognizer.start_continuous_recognition()

    while not done:
        pass

    recognizer.stop_continuous_recognition()

    transcript = " ".join(full_text).strip()
    print("✅ FAST transcription complete.")
    return transcript

# -------------------------------
# Generate SOAP note
# -------------------------------
def generate_soap(transcript):
    # Path to Prompts.txt
    prompt_file = os.path.join(os.path.dirname(__file__), "Prompts.txt")

    # Load prompt text
    with open(prompt_file, "r", encoding="utf-8") as f:
        base_prompt = f.read()

    # Insert transcript into template
    final_prompt = base_prompt.replace("{{TRANSCRIPT}}", transcript)

    headers = {"Content-Type": "application/json", "api-key": OPENAI_KEY}
    url = f"{OPENAI_ENDPOINT.rstrip('/')}/openai/deployments/{OPENAI_DEPLOYMENT}/chat/completions?api-version=2025-01-01-preview"

    body = {
        "messages": [{"role": "user", "content": final_prompt}],
        "max_tokens": 1000,   # GPT-4 supports max_tokens
        "temperature": 0.0    # GPT-4 supports temperature
    }

    print("💬 Generating SOAP note via Azure OpenAI (GPT-4)...")
    print("FINAL PROMPT PREVIEW (first 1000 chars):")
    print(final_prompt[:1000])

    response = requests.post(url, headers=headers, json=body, timeout=120)
    print("DEBUG → Status code:", response.status_code)
    print("DEBUG → Response text:", response.text)

    response.raise_for_status()
    soap_text = response.json()["choices"][0]["message"]["content"].strip()
    print("✅ SOAP note generated, length:", len(soap_text))
    return soap_text

# -------------------------------
# Save SOAP note to Word using template
# -------------------------------



def save_to_docx_with_template(soap_text, template_path, patient_name=None, audio_file_path=None):
    doc = Document(template_path)

    # Replace patient name in headers
    if patient_name:
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    if "{PATIENT_NAME}" in run.text:
                        run.text = run.text.replace("{PATIENT_NAME}", patient_name)

    # Insert SOAP note at the top
    body = doc._body
    lines = [line.strip() for line in soap_text.splitlines() if line.strip()]  # remove blank lines
    for line in reversed(lines):  # reverse so first line appears first
        p = doc.add_paragraph()
        # Detect bold lines marked by **...**
        is_bold = line.startswith("**") and line.endswith("**")
        if is_bold:
            line = line[2:-2].strip()  # remove the ** symbols
        run = p.add_run(line)
        if is_bold:
            run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        body._element.insert(0, p._p)  # insert at top

    # Construct output filename
    if audio_file_path:
        base = os.path.splitext(os.path.basename(audio_file_path))[0]
        filename = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    else:
        filename = f"SOAP_Note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    doc.save(filename)
    print("✅ Saved SOAP note:", filename)
    return filename

# -------------------------------
# Full pipeline
# -------------------------------
def audio_to_docx(audio_file_path, template_path, patient_name=None):
    transcript = transcribe_audio(audio_file_path)
    # -------------------------------
    # Save transcript to "processed" folder
    # -------------------------------
    audio_dir = os.path.dirname(audio_file_path)
    processed_dir = os.path.join(audio_dir, "processed")
    base = os.path.splitext(os.path.basename(audio_file_path))[0]
    transcript_path = os.path.join(
        processed_dir,
        f"{base}_TRANSCRIPT_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    )
    with open(transcript_path, "w", encoding="utf-8") as f:
        f.write(transcript)
    print("📝 Transcript saved to processed folder:", transcript_path)
    soap_text = generate_soap(transcript)
    print("SOAP text length:", len(soap_text))
    print("SOAP preview:\n", soap_text[:500])
    
    return save_to_docx_with_template(soap_text, template_path, patient_name, audio_file_path=audio_file_path)

def process_audio_file(audio_path):
    template_path = os.path.join(os.path.dirname(__file__), "0.docx")
    patient = "John Doe"  # or auto-detect in the future

    try:
        result = audio_to_docx(audio_path, template_path, patient_name=patient)
        print("Generated:", result)
        return result
    except Exception as e:
        print("Processing error:", e)
        return None

